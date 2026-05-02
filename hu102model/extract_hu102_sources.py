"""
呼102井参考资料扫描与关键词索引工具
=====================================

本脚本扫描"参考文档/呼102"目录下的所有文档文件（xlsx、docx、pdf等），
提取文本内容并检索与固井顶替效率相关的关键词命中行，
生成：
1. hu102_source_inventory.json — 每个文件的元信息、预览和关键词命中列表
2. hu102_source_keyword_hits.md — 按文件分组的关键词命中索引（Markdown格式）

用途：快速定位呼102井相关资料中包含关键参数（排量、替浆、井径等）的段落，
为建模提供数据溯源支持。
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

# ========== 路径配置 ==========
ROOT = Path(__file__).resolve().parents[2]  # 项目根目录（脚本所在目录的上两级）
SOURCE_DIR = ROOT / "参考文档" / "呼102"     # 呼102参考资料目录
OUT_DIR = Path(__file__).resolve().parent    # 输出目录：与脚本同目录
OUT_JSON = OUT_DIR / "hu102_source_inventory.json"  # 输出：文件清单JSON
OUT_MD = OUT_DIR / "hu102_source_keyword_hits.md"   # 输出：关键词命中索引Markdown

# ========== 关键词列表 ==========
# 与固井顶替效率建模相关的关键词，用于在文档中检索有价值的信息行
KEYWORDS = (
    "呼102",       # 井号
    "呼探2",       # 井号别名
    "顶替效率",     # 核心关注指标
    "有效顶替",     # 顶替效率子概念
    "固井质量",     # CBL/VDL评价
    "合格率",       # CBL合格率
    "返深",         # 水泥返高深度
    "水泥返",       # 水泥返高
    "水泥浆",       # 水泥浆体系
    "尾浆",         # 尾浆（领浆之后注入）
    "领浆",         # 领浆（先注入的水泥浆）
    "隔离液",       # 隔离液
    "平衡液",       # 平衡液/冲洗液
    "排量",         # 注入排量
    "替浆",         # 替浆作业
    "泵压",         # 泵压
    "井径",         # 井径数据
    "井斜",         # 井斜数据
    "套管",         # 套管程序
    "油气层",       # 油气层段
    "目的层",       # 目的层段
)


def compact(value: Any) -> str:
    """
    将任意值压缩为单行字符串。

    将None转为空串，其他值转为字符串后去除多余空白，
    便于后续关键词检索和预览显示。

    Args:
        value: 任意输入值

    Returns:
        压缩后的单行字符串
    """
    text = "" if value is None else str(value)
    return re.sub(r"\s+", " ", text).strip()


def keyword_hits(text: str) -> list[dict[str, str]]:
    """
    在文本中检索关键词命中行。

    逐行扫描文本，对每个非空行检查是否包含KEYWORDS中的关键词，
    返回命中行的行号、匹配关键词和文本内容。

    Args:
        text: 待检索的全文文本

    Returns:
        命中行列表，每项包含：
        - line: 行号
        - keywords: 匹配到的关键词（逗号分隔）
        - text: 行文本内容（截断至500字符）
        最多返回120条命中记录
    """
    lines = [compact(line) for line in text.splitlines()]
    hits: list[dict[str, str]] = []
    for i, line in enumerate(lines, start=1):
        if not line:
            continue
        matched = [kw for kw in KEYWORDS if kw in line]
        if matched:
            hits.append({"line": str(i), "keywords": ",".join(matched), "text": line[:500]})
    return hits[:120]


def read_xlsx(path: Path) -> dict[str, Any]:
    """
    读取Excel文件，提取工作表信息和关键词命中。

    使用openpyxl以只读模式打开xlsx/xlsm文件，
    读取每个工作表的前80行数据作为预览，
    并在全部提取文本中检索关键词。

    Args:
        path: Excel文件路径

    Returns:
        包含以下字段的字典：
        - kind: "xlsx"
        - sheets: 各工作表的名称、行列数和预览行
        - hits: 关键词命中列表
    """
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets: list[dict[str, Any]] = []
    all_text_parts: list[str] = []
    for ws in wb.worksheets:
        rows: list[list[str]] = []
        # 读取前80行（避免大文件耗时过长）
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 80), values_only=True):
            vals = [compact(v) for v in row]
            if any(vals):
                rows.append(vals[:20])  # 每行最多20列
                all_text_parts.append(" | ".join(vals))
        sheets.append({"name": ws.title, "max_row": ws.max_row, "max_column": ws.max_column, "preview_rows": rows[:25]})
    text = "\n".join(all_text_parts)
    return {"kind": "xlsx", "sheets": sheets, "hits": keyword_hits(text)}


def read_docx(path: Path) -> dict[str, Any]:
    """
    读取Word文档，提取段落、表格和关键词命中。

    使用python-docx打开docx文件，
    提取所有段落的文本和前20个表格的内容，
    并在全文中检索关键词。

    Args:
        path: Word文档路径

    Returns:
        包含以下字段的字典：
        - kind: "docx"
        - char_count: 全文字符数
        - preview: 前80个非空段落的文本
        - hits: 关键词命中列表
    """
    from docx import Document

    doc = Document(str(path))
    # 提取段落文本
    parts = [compact(p.text) for p in doc.paragraphs if compact(p.text)]
    # 提取表格内容（最多20个表格）
    for table in doc.tables[:20]:
        for row in table.rows:
            vals = [compact(cell.text) for cell in row.cells]
            if any(vals):
                parts.append(" | ".join(vals))
    text = "\n".join(parts)
    return {"kind": "docx", "char_count": len(text), "preview": parts[:80], "hits": keyword_hits(text)}


def read_pdf(path: Path) -> dict[str, Any]:
    """
    读取PDF文件，提取文本和关键词命中。

    使用pdfplumber打开PDF文件，
    提取前12页的文本内容作为预览，
    并在提取文本中检索关键词。

    Args:
        path: PDF文件路径

    Returns:
        包含以下字段的字典：
        - kind: "pdf"
        - page_count: 总页数
        - char_count_first_pages: 前几页字符数
        - preview: 前3000字符预览
        - hits: 关键词命中列表
        如果读取失败，返回kind="pdf"和error字段
    """
    text_parts: list[str] = []
    page_count = 0
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            # 只提取前12页（避免大文件耗时过长）
            for page in pdf.pages[:12]:
                text_parts.append(page.extract_text() or "")
    except Exception as exc:  # noqa: BLE001
        return {"kind": "pdf", "error": f"{type(exc).__name__}: {exc}"}
    text = "\n".join(text_parts)
    return {"kind": "pdf", "page_count": page_count, "char_count_first_pages": len(text), "preview": text[:3000], "hits": keyword_hits(text)}


def inspect_file(path: Path) -> dict[str, Any]:
    """
    检查单个文件，提取元信息和内容摘要。

    根据文件扩展名选择对应的读取函数，
    对不支持的格式仅记录基本信息。

    Args:
        path: 文件路径

    Returns:
        文件信息字典，包含：
        - path: 绝对路径
        - relative_path: 相对于项目根目录的路径
        - suffix: 文件扩展名
        - size_bytes: 文件大小
        - kind: 文件类型标识
        - 其他类型特定字段（sheets/preview/hits等）
    """
    suffix = path.suffix.lower()
    rel = str(path.relative_to(ROOT))
    result: dict[str, Any] = {"path": str(path), "relative_path": rel, "suffix": suffix, "size_bytes": path.stat().st_size}
    try:
        if suffix in {".xlsx", ".xlsm"}:
            result.update(read_xlsx(path))
        elif suffix == ".docx":
            result.update(read_docx(path))
        elif suffix == ".pdf":
            result.update(read_pdf(path))
        else:
            result["kind"] = "listed_only"  # 不支持的格式，仅列出
    except Exception as exc:  # noqa: BLE001
        result.update({"kind": "error", "error": f"{type(exc).__name__}: {exc}"})
    return result


def main() -> None:
    """
    主入口：扫描参考资料目录，生成文件清单和关键词索引。

    流程：
    1. 递归搜索参考资料目录中的所有文档文件
    2. 逐个检查文件，提取元信息和关键词命中
    3. 将结果保存为JSON和Markdown格式
    """
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    # 支持的文件格式
    patterns = ("*.xlsx", "*.xlsm", "*.docx", "*.pdf", "*.PDF", "*.jpg", "*.jpeg", "*.png")
    files: list[Path] = []
    for pattern in patterns:
        files.extend(SOURCE_DIR.rglob(pattern))
    files = sorted(set(files), key=lambda p: str(p))  # 去重并排序

    # 逐个检查文件
    inventory = [inspect_file(path) for path in files]

    # 保存JSON格式的完整清单
    OUT_JSON.write_text(json.dumps(inventory, ensure_ascii=False, indent=2), encoding="utf-8")

    # 生成Markdown格式的关键词命中索引
    lines = ["# 呼102资料关键词命中索引", ""]
    for item in inventory:
        hits = item.get("hits") or []
        if not hits:
            continue
        lines.append(f"## {item['relative_path']}")
        lines.append("")
        for hit in hits[:30]:  # 每个文件最多30条命中
            lines.append(f"- L{hit['line']} [{hit['keywords']}]: {hit['text']}")
        lines.append("")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")

    # 输出统计信息
    print(json.dumps({"files": len(inventory), "output_json": str(OUT_JSON), "output_md": str(OUT_MD)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
