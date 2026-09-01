import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
for f in [r"D:\users\desktop\research\控压固井项目\0708\甲方完成数据\HT1-004井油层尾管上交资料1\HT1-004井油层尾管上交资料（甲方）1\化验报告\HT1-004 油层尾管 化验报告.docx",
          r"D:\users\desktop\research\控压固井项目\0708\甲方完成数据\HT1-004井油层尾管上交资料1\HT1-004井油层尾管上交资料（甲方）1\HT1-004 完井尾管 工程量清单.docx"]:
    print("="*90); print("FILE:", f.split("\\")[-1])
    doc = Document(f)
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t: print("P:", t[:300])
    for ti, tbl in enumerate(doc.tables):
        print(f"--- TABLE {ti} ({len(tbl.rows)}x{len(tbl.columns)})")
        for r in tbl.rows:
            cells = [c.text.strip().replace("\n","⏎") for c in r.cells]
            # dedupe merged cells
            out=[]; 
            for c in cells:
                if not out or out[-1]!=c: out.append(c)
            print(" | ".join(out)[:400])
