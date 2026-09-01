import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
base = r"D:\users\desktop\research\控压固井项目\cement model\参考文档\呼1-004"
for sub, f in [("优化参数","优化参数.docx"), ("甲方参数","甲方参数.docx"), ("甲方参数","基础参数代码.docx"), ("设计参数","设计参数.docx")]:
    print("="*90); print("FILE:", sub, f)
    doc = Document(f"{base}\{sub}\{f}")
    for p in doc.paragraphs:
        t = p.text.strip()
        if t: print("P:", t[:260])
    for ti, tbl in enumerate(doc.tables):
        print(f"--- TABLE {ti}")
        for r in tbl.rows:
            cells=[]; 
            for c in r.cells:
                v=c.text.strip().replace("\n","⏎")
                if not cells or cells[-1]!=v: cells.append(v)
            print(" | ".join(cells)[:380])
