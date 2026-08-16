# -*- coding: utf-8 -*-
from docx import Document
import sys
p = r"D:\users\desktop\research\控压固井项目\项目相关数据资料（未重新命名）\甲方完成数据\HT1-004井油层尾管上交资料1\HT1-004井油层尾管上交资料（甲方）1\化验报告\HT1-004 油层尾管 化验报告.docx"
doc = Document(p)
out = []
# paragraphs
for para in doc.paragraphs:
    t = para.text.strip()
    if t:
        out.append(t)
# tables
for ti, table in enumerate(doc.tables):
    out.append(f"===TABLE {ti} rows={len(table.rows)}===")
    for row in table.rows:
        cells = [c.text.strip().replace('\n',' ') for c in row.cells]
        out.append(" | ".join(cells))
txt = "\n".join(out)
with open(r"D:\users\desktop\research\控压固井项目\cement model\参考文档\现场资料提取\ht1_004_呼1-004\_work\化验报告_utf8.txt","w",encoding="utf-8") as f:
    f.write(txt)
print("LEN:", len(txt))
print(txt[:3000])
