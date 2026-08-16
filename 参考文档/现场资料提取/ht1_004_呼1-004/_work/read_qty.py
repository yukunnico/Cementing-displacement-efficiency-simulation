# -*- coding: utf-8 -*-
from docx import Document
p = r"D:\users\desktop\research\控压固井项目\项目相关数据资料（未重新命名）\甲方完成数据\HT1-004井油层尾管上交资料1\HT1-004井油层尾管上交资料（甲方）1\HT1-004 完井尾管 工程量清单.docx"
doc = Document(p)
out=[]
for para in doc.paragraphs:
    t=para.text.strip()
    if t: out.append(t)
for ti,table in enumerate(doc.tables):
    out.append(f"===TABLE {ti} rows={len(table.rows)}===")
    for row in table.rows:
        cells=[c.text.strip().replace('\n',' ') for c in row.cells]
        out.append(" | ".join(cells))
print("\n".join(out))
